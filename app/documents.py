"""Download seguro, cache temporário e extração de conteúdo de anexos.

A API central baixa e transcreve o conteúdo, mas não toma a decisão final de
aprovar ou reprovar a autorização. Essa decisão auditável continua no ERP.
"""
from __future__ import annotations

from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass
from hashlib import sha256
from io import BytesIO
import mimetypes
from pathlib import PurePath
import re
from threading import RLock
import time
from typing import Any
from urllib.parse import urljoin, urlparse

from docx import Document
from pypdf import PdfReader
import requests

from .errors import IntegrationError
from .fracttal import AttachmentMetadata
from .http import build_session
from .security import host_is_allowed
from .settings import Settings

_REDIRECT_CODES = {301, 302, 303, 307, 308}
_IMAGE_MIMES = {"image/jpeg", "image/png", "image/gif", "image/webp"}
_TEXT_MIMES = {"text/plain", "text/csv"}
_SUPPORTED_MIMES = _IMAGE_MIMES | _TEXT_MIMES | {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_SUCCESS_STATUSES = {"EXTRAIDO", "EXTRAIDO_LOCALMENTE", "OCR_REALIZADO"}


@dataclass(frozen=True)
class DownloadedAttachment:
    metadata: AttachmentMetadata
    content: bytes
    mime_type: str
    sha256_hex: str
    download_status: int = 200
    response_content_type: str = ""
    final_host: str = ""
    redirect_count: int = 0


@dataclass(frozen=True)
class ExtractedAttachment:
    id: int
    id_request: int
    description: str
    nome: str
    mime_type: str
    tamanho_bytes: int
    sha256: str
    texto_extraido: str
    metodo_extracao: str
    status_extracao: str
    paginas_analisadas: int
    avisos: tuple[str, ...]
    download_status: int | None = None
    download_content_type: str = ""
    download_final_host: str = ""
    redirecionamentos: int = 0
    error_type: str = ""
    error_stage: str = ""
    evidencias: tuple[str, ...] = ()
    interpretacao: str = ""

    def to_erp_dict(self) -> dict[str, Any]:
        return {
            "id": self.id,
            "id_request": self.id_request,
            "description": self.description,
            "nome": self.nome,
            "mime_type": self.mime_type,
            "tamanho_bytes": self.tamanho_bytes,
            "sha256": self.sha256,
            "texto_extraido": self.texto_extraido,
            "metodo_extracao": self.metodo_extracao,
            "status_extracao": self.status_extracao,
            "paginas_analisadas": self.paginas_analisadas,
            "avisos": list(self.avisos),
            "download_status": self.download_status,
            "download_content_type": self.download_content_type,
            "download_final_host": self.download_final_host,
            "redirecionamentos": self.redirecionamentos,
            "error_type": self.error_type,
            "error_stage": self.error_stage,
            "evidencias": list(self.evidencias),
            "interpretacao": self.interpretacao,
        }


@dataclass(frozen=True)
class _CachedDownload:
    expires_at: int
    downloaded: DownloadedAttachment


# O arquivo já baixado é mantido apenas em memória durante o mesmo prazo das
# URLs de prévia/download. Isso evita uma segunda consulta ao Fracttal e evita
# depender novamente da URL S3 assinada ao clicar no frontend.
_DOWNLOADED_CACHE: dict[tuple[str, int], _CachedDownload] = {}
_DOWNLOADED_CACHE_LOCK = RLock()
_DOWNLOADED_CACHE_MAX_ITEMS = 500


def _purge_download_cache(now: int) -> None:
    expired = [key for key, value in _DOWNLOADED_CACHE.items() if value.expires_at < now]
    for key in expired:
        _DOWNLOADED_CACHE.pop(key, None)


def cache_downloaded_attachment(
    downloaded: DownloadedAttachment,
    settings: Settings,
    *,
    request_code: str | None = None,
) -> None:
    now = int(time.time())
    expires_at = now + settings.attachment_preview_ttl_seconds
    keys = {
        (str(downloaded.metadata.id_request), int(downloaded.metadata.id)),
        ("*", int(downloaded.metadata.id)),
    }
    if request_code:
        keys.add((str(request_code), int(downloaded.metadata.id)))
    with _DOWNLOADED_CACHE_LOCK:
        _purge_download_cache(now)
        entry = _CachedDownload(expires_at=expires_at, downloaded=downloaded)
        for key in keys:
            _DOWNLOADED_CACHE[key] = entry
        if len(_DOWNLOADED_CACHE) > _DOWNLOADED_CACHE_MAX_ITEMS:
            oldest = sorted(
                _DOWNLOADED_CACHE.items(), key=lambda item: item[1].expires_at
            )
            excess = len(_DOWNLOADED_CACHE) - _DOWNLOADED_CACHE_MAX_ITEMS
            for key, _value in oldest[:excess]:
                _DOWNLOADED_CACHE.pop(key, None)


def alias_cached_attachment(code: str, attachment_id: int, settings: Settings) -> bool:
    """Cria uma chave pelo ``code`` sem duplicar os bytes em memória."""
    now = int(time.time())
    with _DOWNLOADED_CACHE_LOCK:
        _purge_download_cache(now)
        source = _DOWNLOADED_CACHE.get(("*", int(attachment_id)))
        if source is None:
            return False
        expires_at = max(
            source.expires_at,
            now + settings.attachment_preview_ttl_seconds,
        )
        _DOWNLOADED_CACHE[(str(code), int(attachment_id))] = _CachedDownload(
            expires_at=expires_at,
            downloaded=source.downloaded,
        )
        return True


def get_cached_downloaded_attachment(
    code: str, attachment_id: int
) -> DownloadedAttachment | None:
    now = int(time.time())
    with _DOWNLOADED_CACHE_LOCK:
        _purge_download_cache(now)
        entry = _DOWNLOADED_CACHE.get((str(code), int(attachment_id)))
        if entry is None:
            entry = _DOWNLOADED_CACHE.get(("*", int(attachment_id)))
        return entry.downloaded if entry is not None else None


def clear_downloaded_attachment_cache() -> None:
    """Uso restrito a testes e reinicializações controladas."""
    with _DOWNLOADED_CACHE_LOCK:
        _DOWNLOADED_CACHE.clear()


def _safe_filename(name: str, attachment_id: int) -> str:
    clean = PurePath(name or "").name.replace("\x00", "").strip()
    return (clean or f"anexo-{attachment_id}")[:255]


def _sniff_mime(content: bytes, header_mime: str, filename: str) -> str:
    if content.startswith(b"%PDF-"):
        return "application/pdf"
    if content.startswith(b"\xff\xd8\xff"):
        return "image/jpeg"
    if content.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if content.startswith((b"GIF87a", b"GIF89a")):
        return "image/gif"
    if len(content) >= 12 and content[:4] == b"RIFF" and content[8:12] == b"WEBP":
        return "image/webp"
    if content.startswith(b"PK\x03\x04") and filename.lower().endswith(".docx"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    normalized_header = (header_mime or "").split(";", 1)[0].strip().lower()
    if normalized_header in _SUPPORTED_MIMES:
        return normalized_header
    guessed, _ = mimetypes.guess_type(filename)
    return guessed or normalized_header or "application/octet-stream"


def _validate_signed_url(url: str, settings: Settings) -> None:
    parsed = urlparse(url)
    if parsed.scheme != "https" or not host_is_allowed(
        parsed.hostname, settings.allowed_attachment_hosts
    ):
        raise IntegrationError(
            "attachment_url_not_allowed",
            "A URL assinada do anexo não pertence a um host permitido.",
            status_code=422,
            stage="baixar_anexo_fracttal",
        )


def _looks_like_storage_error(content: bytes, content_type: str) -> bool:
    sample = content[:4096].lstrip().lower()
    mime = (content_type or "").split(";", 1)[0].strip().lower()
    if mime in {"text/html", "application/xhtml+xml"}:
        return True
    if sample.startswith((b"<!doctype html", b"<html", b"<?xml", b"<error")):
        return any(
            marker in sample
            for marker in (
                b"accessdenied",
                b"expiredtoken",
                b"request has expired",
                b"signaturedoesnotmatch",
                b"<error",
                b"<html",
            )
        )
    return False


def download_attachment(
    metadata: AttachmentMetadata,
    settings: Settings,
) -> DownloadedAttachment:
    if not metadata.signed_url:
        raise IntegrationError(
            "attachment_without_url",
            "O Fracttal retornou um anexo sem URL de download.",
            status_code=422,
            stage="baixar_anexo_fracttal",
            request_code=str(metadata.id_request),
            endpoint=f"attachment/{metadata.id}",
        )
    current_url = metadata.signed_url
    session = build_session(settings)
    response: requests.Response | None = None
    redirects = 0
    try:
        for _ in range(4):
            _validate_signed_url(current_url, settings)
            try:
                response = session.get(
                    current_url,
                    stream=True,
                    allow_redirects=False,
                    headers={"Accept": "*/*", "Accept-Encoding": "identity"},
                    timeout=(settings.connect_timeout, settings.read_timeout),
                )
            except requests.Timeout as exc:
                raise IntegrationError(
                    "attachment_download_timeout",
                    "O download do anexo excedeu o tempo limite.",
                    stage="baixar_anexo_fracttal",
                    request_code=str(metadata.id_request),
                    endpoint=f"attachment/{metadata.id}",
                ) from exc
            except requests.RequestException as exc:
                raise IntegrationError(
                    "attachment_download_error",
                    "Não foi possível conectar ao armazenamento do anexo.",
                    stage="baixar_anexo_fracttal",
                    request_code=str(metadata.id_request),
                    endpoint=f"attachment/{metadata.id}",
                ) from exc

            if response.status_code in _REDIRECT_CODES:
                location = response.headers.get("Location")
                if not location:
                    raise IntegrationError(
                        "attachment_invalid_redirect",
                        "O download do anexo retornou redirecionamento sem destino.",
                        upstream_status=response.status_code,
                        stage="baixar_anexo_fracttal",
                        request_code=str(metadata.id_request),
                        endpoint=f"attachment/{metadata.id}",
                    )
                current_url = urljoin(current_url, location)
                redirects += 1
                response.close()
                response = None
                continue
            break
        else:
            raise IntegrationError(
                "attachment_too_many_redirects",
                "O download do anexo excedeu o limite de redirecionamentos.",
                stage="baixar_anexo_fracttal",
                request_code=str(metadata.id_request),
                endpoint=f"attachment/{metadata.id}",
            )

        if response is None:
            raise IntegrationError(
                "attachment_download_error",
                "O armazenamento não devolveu uma resposta de download.",
                stage="baixar_anexo_fracttal",
                request_code=str(metadata.id_request),
                endpoint=f"attachment/{metadata.id}",
            )
        if response.status_code in (401, 403):
            raise IntegrationError(
                "attachment_signed_url_expired",
                "A URL assinada do anexo expirou ou foi recusada pelo armazenamento.",
                upstream_status=response.status_code,
                stage="baixar_anexo_fracttal",
                request_code=str(metadata.id_request),
                endpoint=f"attachment/{metadata.id}",
            )
        if response.status_code >= 400:
            raise IntegrationError(
                "attachment_download_upstream_error",
                f"O armazenamento do anexo respondeu HTTP {response.status_code}.",
                upstream_status=response.status_code,
                stage="baixar_anexo_fracttal",
                request_code=str(metadata.id_request),
                endpoint=f"attachment/{metadata.id}",
            )

        final_url = response.url or current_url
        final_host = urlparse(final_url).hostname or urlparse(current_url).hostname or ""
        if not host_is_allowed(final_host, settings.allowed_attachment_hosts):
            raise IntegrationError(
                "attachment_redirect_host_not_allowed",
                "O download foi redirecionado para um host não permitido.",
                stage="baixar_anexo_fracttal",
                request_code=str(metadata.id_request),
                endpoint=f"attachment/{metadata.id}",
            )

        raw_length = response.headers.get("Content-Length")
        if raw_length:
            try:
                if int(raw_length) > settings.max_attachment_bytes:
                    raise IntegrationError(
                        "attachment_too_large",
                        "O anexo excede o limite de tamanho configurado.",
                        status_code=413,
                        upstream_status=response.status_code,
                        stage="baixar_anexo_fracttal",
                        request_code=str(metadata.id_request),
                        endpoint=f"attachment/{metadata.id}",
                    )
            except ValueError:
                pass

        chunks: list[bytes] = []
        size = 0
        digest = sha256()
        for chunk in response.iter_content(chunk_size=64 * 1024):
            if not chunk:
                continue
            size += len(chunk)
            if size > settings.max_attachment_bytes:
                raise IntegrationError(
                    "attachment_too_large",
                    "O anexo excede o limite de tamanho configurado.",
                    status_code=413,
                    upstream_status=response.status_code,
                    stage="baixar_anexo_fracttal",
                    request_code=str(metadata.id_request),
                    endpoint=f"attachment/{metadata.id}",
                )
            digest.update(chunk)
            chunks.append(chunk)
        content = b"".join(chunks)
        response_content_type = response.headers.get("Content-Type", "")
        if not content:
            raise IntegrationError(
                "attachment_empty",
                "O armazenamento respondeu com um arquivo vazio.",
                upstream_status=response.status_code,
                stage="baixar_anexo_fracttal",
                request_code=str(metadata.id_request),
                endpoint=f"attachment/{metadata.id}",
            )
        if _looks_like_storage_error(content, response_content_type):
            raise IntegrationError(
                "attachment_signed_url_expired",
                "O armazenamento devolveu uma página de erro no lugar do arquivo; a URL assinada pode ter expirado.",
                upstream_status=response.status_code,
                stage="baixar_anexo_fracttal",
                request_code=str(metadata.id_request),
                endpoint=f"attachment/{metadata.id}",
            )

        filename = _safe_filename(metadata.description, metadata.id)
        mime_type = _sniff_mime(content, response_content_type, filename)
        return DownloadedAttachment(
            metadata=metadata,
            content=content,
            mime_type=mime_type,
            sha256_hex=digest.hexdigest(),
            download_status=response.status_code,
            response_content_type=response_content_type,
            final_host=final_host,
            redirect_count=redirects,
        )
    finally:
        if response is not None:
            response.close()
        session.close()


def _extract_pdf_text(content: bytes, max_pages: int) -> tuple[str, int, list[str]]:
    warnings: list[str] = []
    try:
        reader = PdfReader(BytesIO(content), strict=False)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                return "", 0, ["PDF protegido por senha ou criptografia não suportada"]
        page_count = min(len(reader.pages), max_pages)
        parts: list[str] = []
        for index in range(page_count):
            try:
                text = reader.pages[index].extract_text() or ""
            except Exception:
                text = ""
                warnings.append(f"Não foi possível extrair diretamente a página {index + 1}")
            if text.strip():
                parts.append(text.strip())
        if len(reader.pages) > max_pages:
            warnings.append(
                f"PDF limitado às primeiras {max_pages} páginas para processamento"
            )
        return "\n\n".join(parts).strip(), page_count, warnings
    except Exception as exc:
        return "", 0, [f"PDF inválido ou não processável: {type(exc).__name__}"]


def _extract_docx_text(content: bytes) -> tuple[str, list[str]]:
    try:
        document = Document(BytesIO(content))
        lines = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    lines.append(" | ".join(values))
        return "\n".join(lines).strip(), []
    except Exception as exc:
        return "", [f"DOCX inválido ou não processável: {type(exc).__name__}"]


def _gemini_error_summary(exc: Exception) -> str:
    status = getattr(exc, "status_code", None) or getattr(exc, "code", None)
    suffix = f" (HTTP {status})" if status else ""
    return f"OCR Gemini falhou: {type(exc).__name__}{suffix}"


def _transcribe_with_gemini(
    content: bytes, mime_type: str, settings: Settings
) -> tuple[str, str]:
    if not settings.ocr_with_gemini:
        return "", "OCR por IA não está habilitado na API Central"
    if not settings.document_gemini_api_key:
        return "", (
            "OCR por IA sem chave disponível: nenhuma variável de documento foi "
            "configurada e nenhuma chave Gemini recente foi recebida pelo mesmo ERP"
        )
    try:
        from google import genai
        from google.genai import types

        client = genai.Client(api_key=settings.document_gemini_api_key)
        prompt = (
            "Atue somente como transcritor OCR. Transcreva fielmente todo o texto "
            "visível no documento, preservando nomes, e-mails, datas, cargos, status, "
            "números de requisição, frases de autorização e locais. Não decida se o "
            "documento é válido, não resuma e não complete informações ausentes. "
            "Retorne apenas o texto transcrito."
        )
        response = client.models.generate_content(
            model=settings.document_gemini_model,
            contents=[
                types.Part.from_text(text=prompt),
                types.Part.from_bytes(data=content, mime_type=mime_type),
            ],
            config=types.GenerateContentConfig(temperature=0.0),
        )
        text = str(getattr(response, "text", "") or "").strip()
        if not text:
            return "", "OCR Gemini respondeu sem texto legível"
        return text, ""
    except Exception as exc:
        return "", _gemini_error_summary(exc)


_AUTH_EVIDENCE_RE = re.compile(
    r"\b(?:de acordo|autorizad[oa]s?|aprovad[oa]s?|pode(?:-se)?\s+dar\s+"
    r"prosseguimento|podem?\s+prosseguir|em conformidade|solicita[cç][aã]o\s+"
    r"aprovada|status\s+closed|closed|assinad[oa]|respons[aá]vel|chaveiro|"
    r"c[oó]pia\s+de\s+chaves?)\b",
    re.IGNORECASE,
)


def _document_evidence(text: str) -> tuple[str, ...]:
    if not text.strip():
        return ()
    candidates = re.split(r"(?<=[.!?])\s+|[\r\n]+", text)
    evidence: list[str] = []
    for candidate in candidates:
        clean = " ".join(candidate.split()).strip()
        if clean and _AUTH_EVIDENCE_RE.search(clean):
            evidence.append(clean[:500])
        if len(evidence) >= 8:
            break
    return tuple(evidence)


def _interpretation_for(status: str, text: str, evidence: tuple[str, ...]) -> str:
    if status in _SUCCESS_STATUSES:
        if evidence:
            return (
                "Conteúdo legível extraído. A triagem encontrou trechos potencialmente "
                "relevantes à autorização; a decisão final permanece no ERP."
            )
        return (
            "Conteúdo legível extraído, sem expressão explícita de autorização "
            "identificada pela triagem; a decisão final permanece no ERP."
        )
    if status == "PARCIAL":
        return (
            "Somente parte do conteúdo pôde ser extraída; não é seguro concluir a "
            "validação documental automaticamente."
        )
    if status == "NAO_LEGIVEL":
        return "O arquivo foi baixado, mas nenhum conteúdo legível pôde ser obtido."
    if status == "FORMATO_NAO_SUPORTADO":
        return "O arquivo foi baixado, porém o formato não possui extração automática."
    if status in {"ERRO_DOWNLOAD", "URL_EXPIRADA"}:
        return "O conteúdo do arquivo não pôde ser obtido para análise."
    return "O documento requer conferência manual."


def extract_attachment(
    downloaded: DownloadedAttachment, settings: Settings
) -> ExtractedAttachment:
    metadata = downloaded.metadata
    filename = _safe_filename(metadata.description, metadata.id)
    mime_type = downloaded.mime_type
    text = ""
    method = "none"
    pages = 0
    warnings: list[str] = []
    status = "VALIDACAO_MANUAL"

    if mime_type == "application/pdf":
        local_text, pages, pdf_warnings = _extract_pdf_text(
            downloaded.content, settings.max_pdf_pages
        )
        warnings.extend(pdf_warnings)
        text = local_text
        if local_text:
            method = "PDF_TEXTO_LOCAL"
            status = "PARCIAL" if pdf_warnings else "EXTRAIDO_LOCALMENTE"
        if len(local_text.strip()) < 40 or pdf_warnings:
            ocr_text, ocr_warning = _transcribe_with_gemini(
                downloaded.content, "application/pdf", settings
            )
            if ocr_text:
                text = ocr_text
                method = "OCR_GEMINI_PDF"
                status = "OCR_REALIZADO"
            elif ocr_warning:
                warnings.append(ocr_warning)
        if not text:
            status = "NAO_LEGIVEL"
    elif mime_type in _IMAGE_MIMES:
        text, ocr_warning = _transcribe_with_gemini(
            downloaded.content, mime_type, settings
        )
        if text:
            method = "OCR_GEMINI_IMAGEM"
            status = "OCR_REALIZADO"
        else:
            status = "NAO_LEGIVEL"
            if ocr_warning:
                warnings.append(ocr_warning)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text, docx_warnings = _extract_docx_text(downloaded.content)
        warnings.extend(docx_warnings)
        method = "DOCX_TEXTO_LOCAL" if text else "none"
        status = "EXTRAIDO_LOCALMENTE" if text else "NAO_LEGIVEL"
    elif mime_type in _TEXT_MIMES:
        try:
            text = downloaded.content.decode("utf-8-sig", errors="replace").strip()
        except Exception:
            text = ""
        method = "TEXTO_LOCAL" if text else "none"
        status = "EXTRAIDO_LOCALMENTE" if text else "NAO_LEGIVEL"
    else:
        status = "FORMATO_NAO_SUPORTADO"
        warnings.append(f"Formato não suportado para extração automática: {mime_type}")

    evidence = _document_evidence(text)
    interpretation = _interpretation_for(status, text, evidence)
    return ExtractedAttachment(
        id=metadata.id,
        id_request=metadata.id_request,
        description=metadata.description,
        nome=filename,
        mime_type=mime_type,
        tamanho_bytes=len(downloaded.content),
        sha256=downloaded.sha256_hex,
        texto_extraido=text[:200_000],
        metodo_extracao=method,
        status_extracao=status,
        paginas_analisadas=pages,
        avisos=tuple(dict.fromkeys(warnings)),
        download_status=downloaded.download_status,
        download_content_type=downloaded.response_content_type,
        download_final_host=downloaded.final_host,
        redirecionamentos=downloaded.redirect_count,
        evidencias=evidence,
        interpretacao=interpretation,
    )


def render_attachment_preview(downloaded: DownloadedAttachment) -> tuple[bytes, str]:
    """Retorna uma imagem exibível pelo frontend para fotos e PDFs."""
    if downloaded.mime_type in _IMAGE_MIMES:
        return downloaded.content, downloaded.mime_type
    if downloaded.mime_type == "application/pdf":
        try:
            import pymupdf

            document = pymupdf.open(stream=downloaded.content, filetype="pdf")
            try:
                if document.page_count < 1:
                    raise ValueError("PDF sem páginas")
                page = document.load_page(0)
                pixmap = page.get_pixmap(
                    matrix=pymupdf.Matrix(1.6, 1.6), alpha=False
                )
                return pixmap.tobytes("png"), "image/png"
            finally:
                document.close()
        except Exception as exc:
            raise IntegrationError(
                "attachment_preview_unavailable",
                "Não foi possível gerar a visualização da primeira página do PDF.",
                status_code=422,
                stage="gerar_previa_anexo",
                request_code=str(downloaded.metadata.id_request),
                endpoint=f"attachment/{downloaded.metadata.id}/preview",
            ) from exc
    raise IntegrationError(
        "attachment_preview_unsupported",
        "O formato deste anexo não possui visualização de imagem disponível.",
        status_code=415,
        stage="gerar_previa_anexo",
        request_code=str(downloaded.metadata.id_request),
        endpoint=f"attachment/{downloaded.metadata.id}/preview",
    )


def process_attachment(
    metadata: AttachmentMetadata, settings: Settings
) -> ExtractedAttachment:
    try:
        downloaded = download_attachment(metadata, settings)
        cache_downloaded_attachment(downloaded, settings)
        return extract_attachment(downloaded, settings)
    except IntegrationError as exc:
        if exc.error_type == "attachment_signed_url_expired":
            status = "URL_EXPIRADA"
        else:
            status = "ERRO_DOWNLOAD"
        warnings = [f"{exc.error_type}: {exc.message}"]
        if exc.upstream_status is not None:
            warnings.append(f"HTTP download: {exc.upstream_status}")
        interpretation = _interpretation_for(status, "", ())
        return ExtractedAttachment(
            id=metadata.id,
            id_request=metadata.id_request,
            description=metadata.description,
            nome=_safe_filename(metadata.description, metadata.id),
            mime_type="application/octet-stream",
            tamanho_bytes=0,
            sha256="",
            texto_extraido="",
            metodo_extracao="none",
            status_extracao=status,
            paginas_analisadas=0,
            avisos=tuple(warnings),
            download_status=exc.upstream_status,
            error_type=exc.error_type,
            error_stage=exc.stage or "baixar_anexo_fracttal",
            interpretacao=interpretation,
        )


def process_attachments(
    attachments: list[AttachmentMetadata], settings: Settings
) -> list[ExtractedAttachment]:
    if not attachments:
        return []
    workers = min(settings.attachment_workers, len(attachments))
    with ThreadPoolExecutor(
        max_workers=workers, thread_name_prefix="fracttal-anexo"
    ) as pool:
        return list(
            pool.map(lambda item: process_attachment(item, settings), attachments)
        )
